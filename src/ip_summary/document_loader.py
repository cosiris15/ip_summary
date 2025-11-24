from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Sequence

import pdfplumber
from docx import Document

from .models import LoadedDocument

SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}


def scan_documents(root: Path) -> List[Path]:
    """
    Recursively scan the folder and return supported contract files.
    """
    files: List[Path] = []
    for path in root.rglob("*"):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return sorted(files)


def load_documents(paths: Sequence[Path]) -> List[LoadedDocument]:
    return [load_document(path) for path in paths]


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".docx":
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif suffix == ".pdf":
        text = _read_pdf(path)
    else:
        raise ValueError(f"Unsupported file type for {path}")

    return LoadedDocument(
        path=path,
        text=text.strip(),
        metadata={"filename": path.name, "relative_path": str(path)},
    )


def _read_pdf(path: Path) -> str:
    contents: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            contents.append(page.extract_text() or "")
    return "\n".join(contents)
